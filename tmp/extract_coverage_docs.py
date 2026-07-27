from pathlib import Path
import json
import re

import pdfplumber
from docx import Document

ROOT = Path("tmp/coverage_docs/وثائق التغطية و جداول المنافع")
OUT = Path("tmp/coverage_extracted")
OUT.mkdir(parents=True, exist_ok=True)


def clean(text):
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text)).strip()


def extract_docx(path: Path):
    doc = Document(path)
    data = {"file": str(path), "type": "docx", "paragraphs": [], "tables": []}
    for p in doc.paragraphs:
        t = clean(p.text)
        if t:
            data["paragraphs"].append(t)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        data["tables"].append(rows)
    return data


def extract_pdf(path: Path):
    data = {"file": str(path), "type": "pdf", "pages": [], "tables": []}
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = clean(page.extract_text() or "")
            if text:
                data["pages"].append({"page": i, "text": text})
            for table in page.extract_tables() or []:
                rows = [[clean(cell) for cell in row] for row in table]
                if rows:
                    data["tables"].append({"page": i, "rows": rows})
    return data


def main():
    summary = []
    for path in sorted(ROOT.iterdir()):
        if path.suffix.lower() not in {".docx", ".pdf"}:
            continue
        try:
            data = extract_docx(path) if path.suffix.lower() == ".docx" else extract_pdf(path)
            out_path = OUT / f"{path.stem}.json"
            out_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            summary.append({
                "file": path.name,
                "type": data["type"],
                "paragraphs_or_pages": len(data.get("paragraphs", data.get("pages", []))),
                "tables": len(data.get("tables", [])),
                "output": str(out_path),
            })
        except Exception as exc:
            summary.append({"file": path.name, "error": str(exc)})
    (OUT / "_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
